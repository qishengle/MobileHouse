#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os

def extract_docx_content(filename):
    try:
        from docx import Document
        
        doc = Document(filename)
        text_content = []
        
        print(f"Extracting content from: {filename}")
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # Save extracted content
        output_file = 'mobile_housing_analysis.txt'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))
        
        print(f"Content extracted successfully!")
        print(f"Total text blocks: {len(text_content)}")
        print(f"Output saved to: {output_file}")
        
        # Show first few paragraphs
        print("\n=== DOCUMENT PREVIEW ===")
        for i, para in enumerate(text_content[:10]):
            print(f"\n{i+1}. {para}")
            if len(para) > 200:
                print("...")
        
        return text_content
        
    except ImportError:
        print("python-docx library not found. Installing...")
        import subprocess
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        print("Please run the script again after installation.")
        return None
        
    except Exception as e:
        print(f"Error extracting content: {e}")
        return None

if __name__ == "__main__":
    filename = "移动住房时代，凤凰来仪- 简体.docx"
    if os.path.exists(filename):
        extract_docx_content(filename)
    else:
        print(f"File not found: {filename}") 